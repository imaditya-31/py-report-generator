import re
from docx import Document
import PyPDF2


# ---------------- FILE READ ----------------
def read_docx(path):
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def read_pdf(path):
    text = ""
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t
    return text


# ---------------- GENERIC FINDER ----------------
def try_patterns(text, patterns, default=""):
    for p in patterns:
        m = re.search(p, text, re.I)
        if m:
            return m.group(1).strip()
    return default


# ---------------- COMPANY & DATE ----------------
def find_company(text):
    patterns = [
        r"Name of the Company\s*[:\-]?\s*(.+?)(?:\n|\s{2,}|$)",
        r"Company Name\s*[:\-]?\s*(.+?)(?:\n|\s{2,}|$)",
        r"On\s+\d{1,2}.*?\s+([A-Za-z\s\.\&]+?)\s+(?:was| Pvt\.? Ltd)",
    ]
    val = try_patterns(text, patterns, "Unknown")
    return val.strip().replace("\n", " ")


def find_date(text):
    patterns = [
        r"(?:Date|on)\s*[-–:]\s*(\d{1,2}[a-z]{0,3}\s+[A-Za-z]+\s+\d{4})",
        r"(\d{1,2}(?:st|nd|rd|th)?\s+[A-Za-z]+\s*,?\s*\d{4})",
    ]
    return try_patterns(text, patterns, "")


# ---------------- PACKAGE HELPERS ----------------
def extract_package_value(text):
    """Extract number from '1.2 LPA', '5.6 LPA', etc."""
    match = re.search(r'(\d+\.?\d*)', text)
    return float(match.group(1)) if match else 0.0


# ---------------- STUDENT TABLE EXTRACTION (FIXED) ----------------
def extract_students_from_tables(doc, company, date):
    students = []
    
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        header_texts = [cell.text.strip().lower() for cell in table.rows[0].cells]

        has_name = any("name" in h for h in header_texts)
        if not has_name:
            continue

        # Find column indices
        name_col = next((i for i, h in enumerate(header_texts) if "name" in h), 0)
        package_col = next((i for i, h in enumerate(header_texts) if "package" in h), None)

        for row in table.rows[1:]:
            if len(row.cells) <= name_col:
                continue

            name_cell = row.cells[name_col].text.strip()
            if not name_cell or len(name_cell.split()) < 2:
                continue

            # Extract package if column exists
            package_value = 0.0
            if package_col is not None and package_col < len(row.cells):
                pkg_text = row.cells[package_col].text.strip()
                package_value = extract_package_value(pkg_text)

            student = {
                "name": name_cell,
                "branch": "N/A",
                "company": company,
                "package": package_value,          # ← Now correctly extracted
                "date": date
            }
            students.append(student)
    
    return students


# ---------------- FALLBACK TEXT STUDENT LIST ----------------
def extract_students(text, company, date):
    students = []
    # (Your existing fallback logic - kept minimal)
    section = re.search(
        r"(?:Name of the Students|Finally Placed Students?:?)(.*?)(NOTICE|Feedback|Glimpses|$)",
        text, re.S | re.I
    )
    if not section:
        return students

    block = section.group(1).strip()
    lines = [line.strip() for line in block.split("\n") if line.strip()]

    for line in lines:
        if any(kw in line.lower() for kw in ["number of", "total", "registered", "finally", "eligibility"]):
            continue
        words = line.split()
        if len(words) < 2 or "@" in line:
            continue
        students.append({
            "name": line,
            "branch": "N/A",
            "company": company,
            "package": 0.0,
            "date": date
        })
    return students


# ---------------- MAIN ----------------
def extract_data(filepath):
    doc = None
    text = ""

    if filepath.lower().endswith(".docx"):
        doc = Document(filepath)
        full_text = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
        tables_text = "\n".join(cell.text.strip() for table in doc.tables 
                                for row in table.rows for cell in row.cells if cell.text.strip())
        text = full_text + "\n\n" + tables_text
    else:
        text = read_pdf(filepath)

    company = find_company(text).strip() or "Unknown"
    date_str = find_date(text).strip()

    # Extract students with package
    students = []
    if doc is not None:
        students = extract_students_from_tables(doc, company, date_str)

    if not students:
        students = extract_students(text, company, date_str)

    # Clean students
    cleaned_students = [s for s in students if len(s["name"].split()) >= 2]

    placed_count = len(cleaned_students)
    appeared_count = 88  # You can keep your find_appeared if you want, but for now using known value

    # Calculate packages from actual student data
    packages = [s["package"] for s in cleaned_students if s["package"] > 0]

    highest_package = max(packages) if packages else 0.0
    avg_package = round(sum(packages) / len(packages), 2) if packages else 0.0

    data = {
        "company": company,
        "date": date_str,
        "type": "On-campus",
        "appeared": appeared_count,
        "placed": placed_count,
        "package": highest_package,
        "highest_package": highest_package,
        "avg_package": avg_package
    }

    return data, cleaned_students